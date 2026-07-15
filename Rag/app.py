import os
import hashlib
import tempfile
import streamlit as st
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_groq import ChatGroq
from langchain_core.output_parsers import StrOutputParser
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from langchain_core.documents import Document as LCDocument
from langchain_chroma import Chroma
from langchain_core.runnables import RunnableLambda
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from pypdf import PdfReader
from docx import Document
from dotenv import load_dotenv

st.set_page_config(page_title="Document Q&A (RAG)", page_icon="📄", layout="wide")

load_dotenv()
MAX_MEMORY_LEN = 15


@st.cache_resource(show_spinner=False)
def get_model():
    return ChatGroq(model="llama-3.3-70b-versatile")

@st.cache_resource(show_spinner=False)
def get_embedder():
    return HuggingFaceEmbeddings(model_name="neuml/pubmedbert-base-embeddings")


# Document processing
def extract_text(file_path: str) -> str:
    if file_path.endswith(".docx"):
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)
    elif file_path.endswith(".pdf"):
        reader = PdfReader(file_path)
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    else:
        raise TypeError("Unsupported file type. Upload a .pdf or .docx file.")


def split_text(full_text: str):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500, 
        chunk_overlap=50)
    return splitter.split_documents([LCDocument(page_content=full_text)])


@st.cache_resource(show_spinner="Reading and indexing document...")
def build_vectorstore(file_bytes: bytes, file_suffix: str, file_hash: str):
    """Cached per unique file (via hash) so re-runs don't re-embed."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_suffix) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        full_text = extract_text(tmp_path)
    finally:
        os.remove(tmp_path)

    if not full_text.strip():
        raise ValueError("No extractable text found in this document.")

    chunks = split_text(full_text)
    embedder = get_embedder()

    vectorstore = Chroma.from_documents(
        documents=chunks, 
        embedding=embedder)
    return vectorstore, len(chunks)


def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)


def summarize_memory(model, messages_to_summarize):
    summarize_prompt = ChatPromptTemplate.from_messages([
        ("system", "You are a chat summarizer. Summarize the conversation but keep key facts, ideas, and important details."),
        MessagesPlaceholder(variable_name="memory"),
    ])
    return (summarize_prompt | model).invoke({"memory": messages_to_summarize}).content


# Session state
if "chat_memory" not in st.session_state:
    st.session_state.chat_memory = []  # list of HumanMessage / AIMessage / SystemMessage
if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None
if "chunk_count" not in st.session_state:
    st.session_state.chunk_count = 0



# Sidebar
with st.sidebar:
    st.header("📄 Upload a document")
    uploaded_file = st.file_uploader("PDF or Word (.docx)", type=["pdf", "docx"])

    if uploaded_file is not None:
        file_bytes = uploaded_file.getvalue()
        file_hash = hashlib.md5(file_bytes).hexdigest()
        suffix = os.path.splitext(uploaded_file.name)[1]

        if st.session_state.get("current_file_hash") != file_hash:
            try:
                vectorstore, n_chunks = build_vectorstore(file_bytes, suffix, file_hash)
                st.session_state.vectorstore = vectorstore
                st.session_state.chunk_count = n_chunks
                st.session_state.current_file_hash = file_hash
                st.session_state.chat_memory = []  # reset chat for a new document
                st.success(f"Indexed {uploaded_file.name} ({n_chunks} chunks)")
            except Exception as e:
                st.error(f"Failed to process file: {e}")

    st.divider()
    k = st.slider("Chunks to retrieve (k)", min_value=1, max_value=10, value=3)

    if st.button("🗑️ Clear chat"):
        st.session_state.chat_memory = []
        st.rerun()


# Main chat UI
st.title("📄 Document Q&A (RAG)")
st.caption("Upload a PDF or Word document, then ask questions answered only from its content.")

for msg in st.session_state.chat_memory:
    if isinstance(msg, HumanMessage):
        with st.chat_message("user"):
            st.write(msg.content)
    elif isinstance(msg, AIMessage):
        with st.chat_message("assistant"):
            st.write(msg.content)

user_input = st.chat_input("Ask something about the document...")

if user_input:
    if st.session_state.vectorstore is None:
        st.warning("Please upload a document first.")
    else:
        with st.chat_message("user"):
            st.write(user_input)

        model = get_model()

        
        if len(st.session_state.chat_memory) > MAX_MEMORY_LEN:
            to_keep = st.session_state.chat_memory[-2:]
            to_summarize = st.session_state.chat_memory[:-2]
            summary = summarize_memory(model, to_summarize)
            st.session_state.chat_memory = [SystemMessage(content=f"Summary of earlier conversation: {summary}")] + to_keep

        retriever = st.session_state.vectorstore.as_retriever(search_kwargs={"k": k})

        prompt = ChatPromptTemplate.from_messages([
            ("system", "You are a helpful assistant. Use the following context to answer the user's question:\n{context}\n\n"
             "If the answer isn't in the context, say you don't know. Do not make things up."),
             MessagesPlaceholder(variable_name='chat_history'),
            ("user", "{question}"),
        ])

        chain = (
            {
                "context": RunnableLambda(lambda x: format_docs(retriever.invoke(x["question"]))),
                "chat_history": RunnableLambda(lambda x: x["chat_history"]),
                "question": RunnableLambda(lambda x: x["question"]),
            } | prompt | model | StrOutputParser()
        )

        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                try:
                    response = chain.invoke({
                        "question": user_input,
                        "chat_history": st.session_state.chat_memory,
                    })
                except Exception as e:
                    response = f"An error occurred while processing: {e}"
            st.write(response)

        st.session_state.chat_memory.append(HumanMessage(content=user_input))
        st.session_state.chat_memory.append(AIMessage(content=response))